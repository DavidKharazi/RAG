from datetime import datetime
from typing import Optional, Dict, Any, List, Tuple
from urllib import request
from langchain.retrievers import SelfQueryRetriever
from langchain.chains.query_constructor.base import AttributeInfo
from langchain.retrievers.multi_query import MultiQueryRetriever
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_community.vectorstores import Chroma
from langchain_openai import ChatOpenAI, OpenAIEmbeddings, OpenAI
import sqlite3
from langchain_core.runnables.history import RunnableWithMessageHistory
from langchain_community.chat_message_histories import ChatMessageHistory
from langchain_core.prompts.chat import ChatPromptTemplate, MessagesPlaceholder
from langchain.schema import AIMessage, HumanMessage
from fastapi import FastAPI, WebSocket, WebSocketDisconnect, HTTPException, Form, Request, status, Depends
from fastapi.responses import HTMLResponse, RedirectResponse, FileResponse, JSONResponse
from pydantic import BaseModel
from starlette.staticfiles import StaticFiles
from passlib.context import CryptContext
import uvicorn
import boto3
from fnmatch import fnmatchcase
import json
import docx
import os
from docx.oxml.ns import qn
from docx import Document as DocxDocument
from io import BytesIO
from authlib.integrations.starlette_client import OAuth
from dotenv import load_dotenv
import random
import string
from fastapi.security.oauth2 import OAuth2AuthorizationCodeBearer
from starlette.middleware.sessions import SessionMiddleware
import uuid
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart


load_dotenv()


os.environ['OPENAI_API_KEY'] = 'my_api_key'



model_name = "gpt-4o"
temperature = 0
llm = ChatOpenAI(model=model_name, temperature=temperature)
embeddings = OpenAIEmbeddings()

current_user = 'A100'

# Настройка клиента для Yandex S3
session = boto3.session.Session()
s3_client = session.client(
    service_name='s3',
    endpoint_url='https://storage.yandexcloud.net',
    aws_access_key_id='my_aws',
    aws_secret_access_key='my_secret',
)

CHROMA_PATH = f'./chroma/{current_user}/'

oauth = OAuth()


def init_metadata_db():
    with sqlite3.connect('metadata.db') as conn:
        conn.execute('''
            CREATE TABLE IF NOT EXISTS Admin (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                email TEXT UNIQUE NOT NULL,
                password TEXT NOT NULL
            )
            ''')
        conn.execute('''
                CREATE TABLE IF NOT EXISTS uploaded_docs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                global_source TEXT,
                filename TEXT
                );
                ''')
        conn.execute('''
        CREATE TABLE IF NOT EXISTS Users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            email VARCHAR(255),
            hashed_password VARCHAR(255),
            cyberman_id INTEGER,
            chat_id INTEGER,
            is_active BOOLEAN DEFAULT FALSE,
            confirmation_token TEXT,
            reset_token TEXT,
            new_password VARCHAR(255),
            FOREIGN KEY (cyberman_id) REFERENCES Cyberman(id),
            FOREIGN KEY (chat_id) REFERENCES Chat(id)
        );
        ''')
        conn.execute('''
        CREATE TABLE IF NOT EXISTS Cyberman (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name VARCHAR(255),
            creativity DOUBLE,
            prompt VARCHAR(255)
        );
        ''')
        conn.execute('''
        CREATE TABLE IF NOT EXISTS Session (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            started_at TIMESTAMP,
            ended_at TIMESTAMP,
            user_id INTEGER,
            cyberman_id INTEGER,
            topic TEXT,
            FOREIGN KEY (user_id) REFERENCES Users(id),
            FOREIGN KEY (cyberman_id) REFERENCES Cyberman(id)
        );
        ''')
        conn.execute('''
        CREATE TABLE IF NOT EXISTS Chat (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            messages VARCHAR(255),
            sender VARCHAR(255),
            sent_at TIMESTAMP DEFAULT (datetime('now', 'localtime')),
            session_id INTEGER,
            FOREIGN KEY (session_id) REFERENCES Session(id)
        );
        ''')


init_metadata_db()


class DatabaseManager:
    def __init__(self, db_path="metadata.db"):
        self.connection = sqlite3.connect(db_path)
        self.connection.row_factory = sqlite3.Row  # Позволяет обращаться к колонкам по именам
        self.cursor = self.connection.cursor()
        self.db_path = db_path
        self.connection.row_factory = sqlite3.Row  # Позволяет работать с результатами в виде объектов Row

    def get_chats_by_user_id(self, user_id):
        query = """
        SELECT id, cyberman_id, started_at FROM Session WHERE user_id = ?
        """
        cursor = self.connection.cursor()
        cursor.execute(query, (user_id,))
        rows = cursor.fetchall()
        return [dict(row) for row in rows]

    def create_new_chat_session(self, user_id, cyberman_id):
        query = """
            INSERT INTO Session (user_id, cyberman_id) VALUES (?, ?)
            """
        cursor = self.connection.cursor()
        cursor.execute(query, (user_id, cyberman_id))
        self.connection.commit()
        return cursor.lastrowid

    def get_chat_messages_by_session_id(self, session_id):
        query = """
            SELECT sender, messages, sent_at FROM Chat WHERE session_id = ?
            """
        cursor = self.connection.cursor()
        cursor.execute(query, (session_id,))
        rows = cursor.fetchall()
        return [dict(row) for row in rows]

    def get_or_create_user(self, email, hashed_password, cyberman_id):
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT id FROM Users WHERE email = ?", (email,))
            user = cursor.fetchone()
            if user:
                print(f"Existing user found: {user[0]}")
                cursor.execute("UPDATE Users SET cyberman_id = ? WHERE email = ?", (cyberman_id, email))
                conn.commit()
                return user[0]
            else:
                confirmation_token = str(uuid.uuid4())
                cursor.execute(
                    "INSERT INTO Users (email, hashed_password, cyberman_id, is_active, confirmation_token) VALUES (?, ?, ?, ?, ?)",
                    (email, hashed_password, cyberman_id, False, confirmation_token)
                )
                user_id = cursor.lastrowid
                print(f"New user created: {user_id}")
                return user_id, confirmation_token

    def get_or_create_cyberman(self, name, creativity, prompt):
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT id FROM Cyberman WHERE name = ?", (name,))
            cyberman = cursor.fetchone()
            if cyberman:
                print(f"Existing cyberman found: {cyberman[0]}")
                return cyberman[0]
            else:
                cursor.execute("INSERT INTO Cyberman (name, creativity, prompt) VALUES (?, ?, ?)",
                               (current_user, creativity, prompt))
                cyberman_id = cursor.lastrowid
                print(f"New cyberman created: {cyberman_id}")
                return cyberman_id

    def create_session(self, user_id, cyberman_id):
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            started_at = datetime.now()
            cursor.execute("INSERT INTO Session (started_at, user_id, cyberman_id) VALUES (?, ?, ?)",
                           (started_at, user_id, cyberman_id))
            session_id = cursor.lastrowid
            print(f"New session created: {session_id}")
            return session_id

    def end_session(self, session_id):
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            ended_at = datetime.now()
            cursor.execute("UPDATE Session SET ended_at = ? WHERE id = ?", (ended_at, session_id))

    def add_chat_message(self, session_id, message, sender):
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute("INSERT INTO Chat (messages, sender, session_id) VALUES (?, ?, ?)",
                           (message, sender, session_id))

    def get_session_by_user_and_cyberman(self, user_id, cyberman_id):
        query = """
        SELECT id FROM Session WHERE user_id = ? AND cyberman_id = ?
        """
        self.cursor.execute(query, (user_id, cyberman_id))
        return self.cursor.fetchone()

    def delete_chat(self, session_id: int):
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            # Удаляем сообщения чата
            cursor.execute("DELETE FROM Chat WHERE session_id = ?", (session_id,))
            # Удаляем саму сессию
            cursor.execute("DELETE FROM Session WHERE id = ?", (session_id,))
            conn.commit()


db_manager = DatabaseManager()


class SQLiteChatHistory():
    def __init__(self, db_path="metadata.db", user_email=None, user_password=None, cyberman_name=None):
        self.db_path = db_path
        self.db_manager = DatabaseManager(db_path)
        self.current_session_id = None
        self.user_email = user_email
        self.user_password = user_password
        self.cyberman_name = cyberman_name

    def start_new_session(self, user_email=None, user_password=None, cyberman_name=current_user):
        user_email = user_email or self.user_email
        user_password = user_password or self.user_password
        cyberman_name = cyberman_name or self.cyberman_name

        print(f"Starting new session with: email={user_email}, cyberman={cyberman_name}")

        # Получаем или создаем Cyberman и получаем его ID
        cyberman_id = self.db_manager.get_or_create_cyberman(cyberman_name, temperature, prompt_sys)
        print(f"Cyberman ID: {cyberman_id}")

        # Передаем cyberman_id в метод get_or_create_user
        user_id = self.db_manager.get_or_create_user(user_email, user_password, cyberman_id)
        print(f"User ID: {user_id}")

        # Проверяем, существует ли уже сессия для данного пользователя и Cyberman
        session = self.db_manager.get_session_by_user_and_cyberman(user_id, cyberman_id)
        if session:
            self.current_session_id = session[0]
        else:
            self.current_session_id = self.db_manager.create_session(user_id, cyberman_id)
        print(f"Session ID: {self.current_session_id}")

        return self.current_session_id

    def add_message(self, message):
        if not self.current_session_id:
            print("No active session. Starting a new one.")
            self.start_new_session()

        print(f"Adding message to session {self.current_session_id}")

        # if isinstance(message, HumanMessage):
        #     sender = "human"
        #     content = message.content
        # elif isinstance(message, AIMessage):
        #     sender = "ai"
        #     content = message.content
        # else:
        #     raise ValueError("Invalid message type")
        #
        # self.db_manager.add_chat_message(self.current_session_id, content, sender)

    def messages(self, limit=15):
        if not self.current_session_id:
            return ChatMessageHistory()

        conn = sqlite3.connect(self.db_path)
        c = conn.cursor()
        c.execute(f"SELECT * FROM Chat WHERE session_id = ? ORDER BY id DESC LIMIT ?",
                  (self.current_session_id, limit))
        resp = c.fetchall()[::-1]
        chat_history = []
        for row in resp:
            id, message, sender, sent_at, session_id = row
            if sender == "human":
                chat_history.append(HumanMessage(content=message))
            elif sender == "ai":
                chat_history.append(AIMessage(content=message))
        conn.close()
        return ChatMessageHistory(messages=chat_history)

    def end_session(self):
        if self.current_session_id:
            self.db_manager.end_session(self.current_session_id)
            self.current_session_id = None


chat_history = SQLiteChatHistory()

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")


def get_password_hash(password):
    return pwd_context.hash(password)


def add_user_to_db(email: str, password: str = None, cyberman_id: int = None, chat_id: int = None):
    hashed_password = get_password_hash(password) if password else None
    confirmation_token = str(uuid.uuid4())  # Генерация токена подтверждения

    try:
        with sqlite3.connect('metadata.db') as conn:
            cursor = conn.cursor()
            cursor.execute(
                "INSERT INTO Users (email, hashed_password, cyberman_id, chat_id, is_active, confirmation_token) VALUES (?, ?, ?, ?, ?, ?)",
                (email, hashed_password, cyberman_id, chat_id, False, confirmation_token)
            )
            conn.commit()
            user_id = cursor.lastrowid  # Получаем ID добавленного пользователя
            return user_id, confirmation_token

    except sqlite3.IntegrityError:
        raise HTTPException(status_code=400, detail="User already registered")


def authenticate_user(email: str, password: str):
    with sqlite3.connect('metadata.db') as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT hashed_password, is_active FROM users WHERE email = ?", (email,))
        row = cursor.fetchone()

    if row:
        hashed_password, is_active = row
        if is_active and pwd_context.verify(password, hashed_password):
            return True
    return False


def delete_chat_history_last_n(self, n=10):
    conn = sqlite3.connect(self.db_path)
    c = conn.cursor()
    c.execute(f'''
    with max_id as (select max(id) as maxid from history_messages where user_id = '{current_user}')
    DELETE FROM history_messages
    WHERE id BETWEEN (select maxid from max_id) - {n} AND (select maxid from max_id)
    ''')
    conn.commit()
    conn.close()


def add_filename_to_metadata(source, filename):
    with sqlite3.connect('metadata.db') as conn:
        conn.execute(f'''INSERT INTO uploaded_docs (global_source, filename) values ('{source}', '{filename}') ; ''')


def delete_filename_from_metadata(source, filename):
    with sqlite3.connect('metadata.db') as conn:
        conn.execute(f'''DELETE from uploaded_docs where global_source = '{source}' and filename ='{filename}' ; ''')


class Document:
    def __init__(self, source: str, page_content: str, metadata: Optional[Dict[str, Any]] = None):
        self.source = source
        self.page_content = page_content
        self.metadata = metadata if metadata is not None else {'source': source}


def get_uploaded_filenames(source) -> List[str]:
    with sqlite3.connect('metadata.db') as conn:
        cursor = conn.cursor()
        cursor.execute(f"SELECT filename FROM uploaded_docs WHERE global_source = ?", (source,))
        rows = cursor.fetchall()
    filenames = [row[0] for row in rows]
    return filenames


def load_s3_files(bucket: str, prefix: str, suffix: str) -> List[str]:
    """List files in a given S3 bucket with a specified prefix and suffix."""
    try:
        response = s3_client.list_objects_v2(Bucket=bucket, Prefix=prefix)
        files = [content['Key'] for content in response.get('Contents', []) if content['Key'].endswith(suffix)]
        if not files:
            print(f"No files found in bucket {bucket} with prefix {prefix} and suffix {suffix}")
        else:
            print(f"Files found in bucket {bucket} with prefix {prefix} and suffix {suffix}: {files}")
        return files
    except Exception as e:
        print(f"Error listing files in bucket {bucket} with prefix {prefix} and suffix {suffix}: {e}")
        return []


def load_docx_new(source, bucket: str) -> List[Document]:
    prefix = 'A100/docx/'
    suffix = '.docx'
    files = load_s3_files(bucket, prefix, suffix)
    uniq_files = get_uploaded_filenames(source) or []

    docs = []
    for file in files:
        if not any(fnmatchcase(file, f"*{pattern}*") for pattern in uniq_files):
            try:
                obj = s3_client.get_object(Bucket=bucket, Key=file)
                content = obj['Body'].read()

                # Используем BytesIO для чтения содержимого файла как бинарного потока
                doc_stream = BytesIO(content)
                doc = DocxDocument(doc_stream)

                # Извлекаем текст из документа docx
                full_text = []
                image_counter = 1

                # Получаем имя файла без расширения и создаем соответствующую папку
                filename_without_extension = os.path.splitext(os.path.basename(file))[0]
                image_folder = filename_without_extension  # Используем оригинальное имя файла для папки

                for para in doc.paragraphs:
                    # Обработка параграфов для создания ссылок на изображения
                    para_text = para.text
                    for run in para.runs:
                        for drawing in run.element.findall('.//a:blip', namespaces={
                            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}):
                            image_rId = drawing.get(qn('r:embed'))
                            image_part = doc.part.related_parts[image_rId]
                            image_filename = f'image_{image_counter:02d}.{image_part.content_type.split("/")[-1]}'
                            image_counter += 1

                            # Загрузка изображения в бакет Яндекса
                            img_content = image_part.blob
                            s3_image_key = f"A100/images/{image_folder}/{image_filename}"
                            s3_client.put_object(
                                Bucket=bucket,
                                Key=s3_image_key,
                                Body=img_content,
                                ContentDisposition='inline',
                                ContentType=image_part.content_type
                            )

                            # Генерация URL для изображения
                            s3_image_url = f"https://storage.yandexcloud.net/{bucket}/{s3_image_key}"
                            para_text += f'\n{s3_image_url}'
                    full_text.append(para_text)
                content = '\n'.join(full_text)

                docs.append(Document(source=file, page_content=content))
            except Exception as e:
                print(f"Error reading docx file {file}: {e}")

    return docs if docs else None


def load_txts(source, bucket: str) -> List[Document]:
    prefix = f'{current_user}/txt/'
    suffix = '.txt'
    files = load_s3_files(bucket, prefix, suffix)
    uniq_files = get_uploaded_filenames(source) or []

    docs = []
    for file in files:
        if not any(fnmatchcase(file, f"*{pattern}*") for pattern in uniq_files):
            try:
                obj = s3_client.get_object(Bucket=bucket, Key=file)
                content = obj['Body'].read().decode('utf-8')
                docs.append(Document(source=file, page_content=content))
            except Exception as e:
                print(f"Error reading txt file {file}: {e}")

    return docs if docs else None


def load_jsons(source, bucket: str) -> Tuple[List[Document], List[dict]]:
    prefix = f'{current_user}/json/'
    suffix = '.json'
    files = load_s3_files(bucket, prefix, suffix)
    uniq_files = get_uploaded_filenames(source) or []

    json_docs, json_metadata = [], []
    for file in files:
        if not any(fnmatchcase(file, f"*{pattern}*") for pattern in uniq_files):
            try:
                obj = s3_client.get_object(Bucket=bucket, Key=file)
                content = json.loads(obj['Body'].read().decode('utf-8'))
                json_docs.append(content)
                json_metadata.append({'source': file})
            except Exception as e:
                print(f"Error reading json file {file}: {e}")

    return (json_docs, json_metadata) if json_docs else (None, None)


def load_documents(global_source, bucket: str, file_types: List[str]) -> dict:
    """
    Загружаем документы в зависимости от типа документа из Yandex S3
    """
    all_docs = {'txt': None, 'json': None, 'json_metadata': None, 'docx': None}
    if 'txt' in file_types:
        txt_docs = load_txts(global_source, bucket)
        all_docs['txt'] = txt_docs
    if 'json' in file_types:
        json_docs, json_metadata = load_jsons(global_source, bucket)
        all_docs['json'] = json_docs
        all_docs['json_metadata'] = json_metadata
    if 'docx' in file_types:
        docx_docs = load_docx_new(global_source, bucket)
        all_docs['docx'] = docx_docs
    return all_docs


# Пример использования
DATA_BUCKET = 'utlik'
DOCS = load_documents('s3', DATA_BUCKET, ['txt', 'json', 'docx'])


def split_docs_to_chunks(documents: dict, file_types: List[str], chunk_size=2000, chunk_overlap=500):
    all_chunks = []
    if 'txt' in file_types and documents['txt'] is not None:
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        txt_chunks = [text_splitter.split_documents([doc]) for doc in documents['txt']]
        txt_chunks = [item for sublist in txt_chunks for item in sublist]
        all_chunks.extend(txt_chunks)

    if 'json' in file_types and documents['json'] is not None:
        json_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        json_chunks = json_splitter.create_documents([json.dumps(doc, ensure_ascii=False) for doc in documents['json']],
                                                     metadatas=documents['json_metadata'])
        all_chunks.extend(json_chunks)

    if 'docx' in file_types and documents['docx'] is not None:
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        txt_chunks = [text_splitter.split_documents([doc]) for doc in documents['docx']]
        txt_chunks = [item for sublist in txt_chunks for item in sublist]
        all_chunks.extend(txt_chunks)

    return all_chunks


chunks_res = split_docs_to_chunks(DOCS, ['txt', 'json', 'docx'])


def get_chroma_vectorstore(documents, embeddings, persist_directory):
    if os.path.isdir(persist_directory) and os.listdir(persist_directory):
        print("Loading existing Chroma vectorstore...")
        vectorstore = Chroma(
            embedding_function=embeddings, persist_directory=persist_directory
        )

        existing_files = get_uploaded_filenames('local')
        uniq_sources_to_add = set(
            doc.metadata['source'] for doc in chunks_res
            if doc.metadata['source'] not in existing_files
        )

        if uniq_sources_to_add:
            vectorstore.add_documents(
                documents=[doc for doc in chunks_res if doc.metadata['source'] in uniq_sources_to_add],
                embedding=embeddings
            )
            for filename in uniq_sources_to_add:
                add_filename_to_metadata('local', filename)
        else:
            print('Новых документов не было, пропускаем шаг добавления')

    else:
        print("Creating and indexing new Chroma vectorstore...")
        vectorstore = Chroma.from_documents(
            documents=documents,
            embedding=embeddings, persist_directory=persist_directory
        )
        uniq_sources_to_add = set(doc.metadata['source'] for doc in documents)
        for filename in uniq_sources_to_add:
            add_filename_to_metadata('local', filename)

    return vectorstore


vectorstore = get_chroma_vectorstore(documents=chunks_res, embeddings=embeddings, persist_directory=CHROMA_PATH)

retriever = vectorstore.as_retriever(search_kwargs={"k": 2}, search_type='similarity')

# retriever = vectorstore.as_retriever(
#     search_type='similarity_score_threshold',
#     search_kwargs={"k": 3,  "score_threshold": 0.1},
#     )


# retriever = MultiQueryRetriever.from_llm(
#     retriever=vectorstore.as_retriever(), llm=llm
# )
#
# # Set logging for the queries
# import logging
#
# logging.basicConfig()
# logging.getLogger("langchain.retrievers.multi_query").setLevel(logging.INFO)




# metadata_field_info = [
#     AttributeInfo(
#         name="категория",
#         description="Категория документа, например, 'ЭЦП', 'документооборот'",
#         type="string",
#     ),
#     AttributeInfo(
#         name="действие",
#         description="Тип действия, например, 'отклонение', 'подписание', 'отмена'",
#         type="string",
#     ),
# ]
#
# document_content_description = "Инструкции по работе с электронной документацией и ЭЦП"
#
# retr = SelfQueryRetriever.from_llm(
#     llm,
#     vectorstore,
#     document_content_description,
#     metadata_field_info,
#     verbose=True
# )
#
# ret = vectorstore.as_retriever(search_kwargs={"k": 2}, search_type='similarity')
#
#
# from langchain.retrievers import EnsembleRetriever
#
# retriever = EnsembleRetriever(
#     retrievers=[retr, ret],
#     weights=[0.5, 0.5]
# )
#
#
# def post_process_results(results):
#     keywords = ["ЭЦП", "электронная подпись", "отклонение", "отмена"]
#     return [r for r in results if any(keyword in r.page_content for keyword in keywords)]
#
# raw_results = retriever.get_relevant_documents("Как отклонить ЭЦП")
# filtered_results = post_process_results(raw_results)





def format_docs(docs):
    return "\n\n".join(doc.page_content for doc in docs)


chat_history_for_chain = SQLiteChatHistory()



# prompt_sys = '''
#         Context:
#             {context}
#         You are an assistant for question-answering tasks. Your aim is to help to new employees with job-specific questions.
#         You should help them with propper actions, recomendations abouts software using the following pieces of retrieved context.
#         You can also use information from chat_history to better understand the problem if necessary.
#         Use only {context} for consultation. Do not search for information on the Internet
#         First understand the meaning of the user's question, and then look for information in {context}.
#         If you don't find the answer in the {context}, just say 'I don't know', e.g.:
#         Answer the question based only on the context above. If the answer is not in the context, say "Из представленного контекста ответа нет".
#         If you meet links to the images in your context always display them in your response.
#         The context which you should use: {context}
#         Question: {question}
# '''



# prompt_sys = '''
#         Context:
#             {context}
#         You are an assistant for question-answering tasks, aimed at helping new employees with job-specific questions.
#         Your goal is to provide proper actions and recommendations regarding software based on the retrieved context.
#         Strictly use only the information provided in the {context}. Do not search for additional information online, tell jokes, or discuss topics unrelated to the given context. If the information is not in the context, simply state: "Из представленного контекста ответа нет".
#         Use only the provided {context} for consultation. Do not search for information on the Internet or outside the given context.
#         First, understand the user's question, and then look for information in the {context}.
#         If the answer is not directly in the {context}, analyze the context again, looking for semantic similarities with the query. If after this analysis the answer is still not found, respond: "Ответ не найден, пожалуйста уточните ваш вопрос".
#         If the context includes links to images, display them in your response.
#         Context: {context}
#         Question: {question}
# '''


prompt_sys = '''
Вы - ассистент для ответов на вопросы, предназначенный для помощи новым сотрудникам с вопросами, связанными с работой.
Ваша цель - предоставлять правильные действия и рекомендации относительно программного обеспечения на основе предоставленного контекста.

ВАЖНО: Вы АБСОЛЮТНО ОГРАНИЧЕНЫ использованием ТОЛЬКО информации, предоставленной в следующем контексте:

Контекст:
{context}

СТРОГИЕ ПРАВИЛА:
1. Используйте ИСКЛЮЧИТЕЛЬНО информацию из предоставленного контекста. НИКОГДА не обращайтесь к внешним источникам или своим знаниям.
2. НЕ ВЫДУМЫВАЙТЕ информацию. Если ответа нет в контексте, скажите "Ответ не найден, пожалуйста, уточните ваш вопрос".
3. НЕ рассказывайте анекдоты, НЕ обсуждайте темы, не связанные с контекстом.
4. Если в контексте есть ссылки на изображения, отобразите их в вашем ответе.

ПРОЦЕСС ОТВЕТА:
1. Внимательно прочитайте вопрос пользователя.
2. Проанализируйте предоставленный контекст на наличие прямого ответа.
3. Если прямой ответ не найден, повторно проанализируйте контекст, ища семантические сходства с вопросом.
4. Если ответ все еще не найден, ответьте: "Ответ не найден, пожалуйста, уточните ваш вопрос".

Вопрос: {question}
'''




prompt_new = ChatPromptTemplate.from_messages(
    [
        (
            "system", prompt_sys,
        ),
        MessagesPlaceholder(variable_name="chat_history"),
        ("human", "{question}"),
    ]
)

chain_new = prompt_new | llm

chain_with_message_history = RunnableWithMessageHistory(
    chain_new,
    lambda session_id: chat_history_for_chain.messages(limit=15),
    input_messages_key="question",
    history_messages_key="chat_history",
)

app = FastAPI()

app.add_middleware(SessionMiddleware, secret_key="1234567890")

oauth = OAuth()


def generate_random_password(length: int = 8):
    characters = string.ascii_letters + string.digits
    return ''.join(random.choice(characters) for _ in range(length))


def send_reset_password_email(email: str, new_password: str):
    # Здесь должна быть логика для отправки email с новым паролем
    print(f"Отправка нового пароля {new_password} на email {email}")


# Подключение статических файлов
app.mount("/static", StaticFiles(directory="static"), name="static")


@app.get("/register", response_class=HTMLResponse)
async def get_register():
    return FileResponse("static/register.html")


def is_email_unique(email: str) -> bool:
    """Проверяет, является ли email уникальным в таблице Users."""
    with sqlite3.connect('metadata.db') as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT 1 FROM Users WHERE email = ?", (email,))
        return cursor.fetchone() is None


# def send_confirmation_email(email: str, token: str):
#     # Настройки для отправки почты
#     sender_email = "datohar82@gmail.com"
#     receiver_email = email
#     password = "my_password"  # Не храните пароли в коде в реальных приложениях
#
#     message = MIMEMultipart("alternative")
#     message["Subject"] = "Подтверждение регистрации"
#     message["From"] = sender_email
#     message["To"] = receiver_email
#
#     confirmation_url = f"http://localhost:8222/confirm-email?token={token}"
#     # confirmation_url = f"https://chata100.up.railway.app/confirm-email?token={token}"
#
#     text = f"Пожалуйста, подтвердите вашу регистрацию, перейдя по следующей ссылке: {confirmation_url}"
#     html = f"""
#     <html>
#     <body>
#         <p>Пожалуйста, подтвердите вашу регистрацию, перейдя по следующей ссылке:<br>
#         <a href="{confirmation_url}">Подтвердить регистрацию</a>
#         </p>
#     </body>
#     </html>
#     """
#
#     part1 = MIMEText(text, "plain")
#     part2 = MIMEText(html, "html")
#
#     message.attach(part1)
#     message.attach(part2)
#
#     try:
#         # Используем SMTP-сервер Gmail
#         server = smtplib.SMTP_SSL("smtp.gmail.com", 465)
#         server.login(sender_email, password)
#         server.sendmail(sender_email, receiver_email, message.as_string())
#         server.quit()
#         print("Письмо успешно отправлено")
#     except Exception as e:
#         print(f"Ошибка при отправке письма: {e}")




def send_confirmation_email(user_email: str, token: str):
    # Настройки для отправки почты
    sender_email = "datohar82@gmail.com"
    intermediate_email = "datohar@icloud.com"
    password = "my_password"  # Не храните пароли в коде в реальных приложениях

    # Создаем сообщение для промежуточного адреса (администратора)
    admin_message = MIMEMultipart("alternative")
    admin_message["Subject"] = f"Подтверждение регистрации для {user_email}"
    admin_message["From"] = sender_email
    admin_message["To"] = intermediate_email

    # confirmation_url = f"http://localhost:8222/confirm-email?token={token}"
    confirmation_url = f"https://chata100.up.railway.app/confirm-email?token={token}"

    admin_text = f"""
    Получена новая регистрация для {user_email}.
    Ссылка подтверждения: {confirmation_url}
    Пожалуйста, перешлите эту ссылку пользователю.
    """
    admin_html = f"""
    <html>
    <body>
        <p>Получена новая регистрация для {user_email}.</p>
        <p>Ссылка подтверждения: <a href="{confirmation_url}">Подтвердить регистрацию</a></p>
        <p>Пожалуйста, перешлите эту ссылку пользователю.</p>
    </body>
    </html>
    """

    admin_part1 = MIMEText(admin_text, "plain")
    admin_part2 = MIMEText(admin_html, "html")

    admin_message.attach(admin_part1)
    admin_message.attach(admin_part2)

    # Сообщение для пользователя
    user_message = MIMEMultipart("alternative")
    user_message["Subject"] = "Ваша заявка принята"
    user_message["From"] = sender_email
    user_message["To"] = user_email

    user_text = f"""
    Уважаемый пользователь,

    Ваша заявка на регистрацию получена и будет рассмотрена администратором в ближайшее время.
    Вы получите уведомление после проверки.

    С уважением,
    Ваша команда поддержки.
    """
    user_html = f"""
    <html>
    <body>
        <p>Уважаемый пользователь,</p>
        <p>Ваша заявка на регистрацию получена и будет рассмотрена администратором в ближайшее время.</p>
        <p>Вы получите уведомление после проверки.</p>
        <p>С уважением,<br>Ваша команда поддержки.</p>
    </body>
    </html>
    """

    user_part1 = MIMEText(user_text, "plain")
    user_part2 = MIMEText(user_html, "html")

    user_message.attach(user_part1)
    user_message.attach(user_part2)

    try:
        # Используем SMTP-сервер Gmail
        server = smtplib.SMTP_SSL("smtp.gmail.com", 465)
        server.login(sender_email, password)

        # Отправляем письмо администратору
        server.sendmail(sender_email, intermediate_email, admin_message.as_string())
        print(f"Письмо успешно отправлено на {intermediate_email}")

        # Отправляем письмо пользователю
        server.sendmail(sender_email, user_email, user_message.as_string())
        print(f"Письмо успешно отправлено пользователю на {user_email}")

        server.quit()
    except Exception as e:
        print(f"Ошибка при отправке письма: {e}")


@app.post("/register")
async def post_register(username: str = Form(...), password: str = Form(...)):
    if not is_email_unique(username):
        return JSONResponse(content={"status": "error", "message": "Пользователь с таким именем уже существует."},
                            status_code=401)
    else:
        user_id, confirmation_token = add_user_to_db(username, password)
        send_confirmation_email(username, confirmation_token)
        return JSONResponse(
            content={"status": "success", "message": "На вашу почту отправлено письмо с подтверждением."},
            status_code=200)


@app.get("/confirm-email")
async def confirm_email(token: str):
    try:
        with sqlite3.connect('metadata.db') as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT id, email FROM Users WHERE confirmation_token = ?", (token,))
            user = cursor.fetchone()

            if user:
                user_id, username = user

                # Обновление статуса пользователя
                cursor.execute(
                    "UPDATE Users SET is_active = TRUE, confirmation_token = NULL WHERE id = ?",
                    (user_id,)
                )
                conn.commit()

                # Логика создания сессии и добавления сообщения
                chat_history = SQLiteChatHistory(user_email=username, user_password="dummy_password")
                print(f"SQLiteChatHistory создан с user_email={username}")
                try:
                    session_id = chat_history.start_new_session()
                    if not check_session_id_exists(session_id):
                        print(f"Session started: {session_id}")
                        db_manager.add_chat_message(session_id,
                                                    "Вас приветствует А100! Напишите Ваш вопрос о документообороте.",
                                                    "Система")
                    else:
                        print(f"Session started: {session_id}")
                    # Переадресация на клиентскую часть после активации
                    return RedirectResponse(url="/")
                except Exception as e:
                    print(f"Error starting session: {e}")
                    return JSONResponse(content={"status": "error", "message": "Failed to start session"},
                                        status_code=500)
            else:
                # return JSONResponse(content={"status": "error", "message": "Invalid or expired token."}, status_code=400)
                return FileResponse("static/error_page.html")

    except Exception as e:
        print(f"Error confirming email: {e}")
        raise HTTPException(status_code=500, detail="Internal Server Error")


@app.get("/check_email")
async def check_email(email: str):
    is_unique = is_email_unique(email)
    return JSONResponse(content={"is_unique": is_unique})


@app.get("/login", response_class=HTMLResponse)
async def get_login():
    return FileResponse("static/login.html")


# Функция для поиска по session_id в Chat
def check_session_id_exists(session_id: str) -> bool:
    with sqlite3.connect('metadata.db') as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT EXISTS(SELECT 1 FROM Chat WHERE session_id = ?)", (session_id,))
        return cursor.fetchone()[0] == 1


@app.post("/login")
async def login(username: str = Form(...), password: str = Form(...)):
    user_email = username
    if authenticate_user(username, password):
        chat_history = SQLiteChatHistory(user_email=user_email, user_password=password)
        print(f"SQLiteChatHistory создан с user_email={user_email}")
        try:
            session_id = chat_history.start_new_session()
            if not check_session_id_exists(session_id):
                print(f"Session started: {session_id}")
                db_manager.add_chat_message(session_id,
                                            "Вас приветствует А100! Напишите Ваш вопрос о документообороте.", "Система")
                return JSONResponse(content={"status": "success", "redirect": "/"})
            else:
                print(f"Session started: {session_id}")
                return JSONResponse(content={"status": "success", "redirect": "/"})
        except Exception as e:
            print(f"Error starting session: {e}")
            return JSONResponse(content={"status": "error", "message": "Failed to start session"}, status_code=500)
    else:
        return JSONResponse(content={"status": "error", "message": "Пользователь не найден или не активирован"},
                            status_code=401)


@app.get("/check_login")
async def check_login(username: str, password: str):
    if authenticate_user(username, password):
        return JSONResponse(content={"is_valid": True})
    else:
        return JSONResponse(content={"is_valid": False})


@app.get('/auth')
async def auth(request: Request):
    token = await oauth.google.authorize_access_token(request)
    user = await oauth.google.parse_id_token(request, token)
    if user:
        add_user_to_db(user['name'], None, user['email'], user['sub'])
        return RedirectResponse(url="/")
    else:
        return HTMLResponse("Authorization failed", status_code=400)


@app.get('/login/google')
async def login_google(request: Request):
    redirect_uri = 'http://localhost:8222/auth'
    return await oauth.google.authorize_redirect(request, redirect_uri)


@app.get("/", response_class=HTMLResponse)
async def read_root():
    return FileResponse("static/index.html")


# Функция для поиска пользователя по email
def get_user_by_email(email: str):
    with sqlite3.connect('metadata.db') as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id, cyberman_id FROM Users WHERE email = ?", (email,))
        return cursor.fetchone()


# Функция для поиска сессии по user_id и cyberman_id
def get_session_by_user_and_cyberman(user_id: int, cyberman_id: int):
    with sqlite3.connect('metadata.db') as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM Session WHERE user_id = ? AND cyberman_id = ?", (user_id, cyberman_id))
        return cursor.fetchone()


@app.post("/create_new_chat/")
async def create_new_chat(request: Request):

    data = await request.json()
    email = data.get('email')

    if not email:
        raise HTTPException(status_code=400, detail="Email is required")

    user = get_user_by_email(email)
    if user is None:
        raise HTTPException(status_code=404, detail="User not found")

    user_id, cyberman_id = user
    session_id = db_manager.create_new_chat_session(user_id, cyberman_id)

    # Создаем начальное сообщение в чате с использованием нового session_id
    db_manager.add_chat_message(session_id, "Вас приветствует А100! Напишите Ваш вопрос о документообороте.", "Система")

    return {"session_id": session_id}


#Модель для анализа сообщений
# llm = ChatOpenAI(model="gpt-3.5-turbo", temperature=0, base_url=BASE_URL)
llm = ChatOpenAI(model="gpt-3.5-turbo", temperature=0)

analyze_prompt = ChatPromptTemplate.from_messages([
    ("system", "Проанализируйте следующие сообщения и сформулируйте тему переписки 2 - 3 словах, не больше 23 символов в сумме."),
    ("human", "{messages}")
])


def analyze_current_chat_topic(user_id, current_session_id):
    with sqlite3.connect('metadata.db') as conn:
        cursor = conn.cursor()

        # Получаем все сообщения текущей сессии
        messages = db_manager.get_chat_messages_by_session_id(current_session_id)

        # Проверяем количество сообщений
        if len(messages) >= 3:
            # Извлекаем второе и третье сообщения
            relevant_messages = messages[1:3]

            # Собираем текст только второго и третьего сообщений
            messages_text = "\n".join([f"{msg['sender']}: {msg['messages']}" for msg in relevant_messages])

            # Анализируем текст сообщений
            chain = analyze_prompt | llm
            topic = chain.invoke({"messages": messages_text}).content

            # Обновляем тему текущей сессии в базе данных
            cursor.execute("UPDATE Session SET topic = ? WHERE id = ?", (topic, current_session_id))
            conn.commit()

def update_session_topic(user_id, session_id, new_topic):
    with sqlite3.connect('metadata.db') as conn:
        cursor = conn.cursor()

        # Обновляем тему в сессии
        cursor.execute("UPDATE Session SET topic = ? WHERE id = ? AND user_id = ?",
                       (new_topic, session_id, user_id))
        conn.commit()


class UpdateTopicRequest(BaseModel):
    session_id: int
    new_topic: str

@app.put("/update-topic")
async def update_topic(request: UpdateTopicRequest):
    try:
        with sqlite3.connect('metadata.db') as conn:
            cursor = conn.cursor()
            cursor.execute(
                "UPDATE Session SET topic = ? WHERE id = ?",
                (request.new_topic, request.session_id)
            )
            if cursor.rowcount == 0:
                raise HTTPException(status_code=404, detail="Session not found or unauthorized")
            conn.commit()
        return {"status": "success", "message": "Topic updated successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/get_chat_messages/{session_id}")
async def get_chat_messages(session_id: int):
    messages = db_manager.get_chat_messages_by_session_id(session_id)
    return {"messages": messages}


@app.get("/get_user_chats/{email}")
async def get_user_chats(email: str):
    user = get_user_by_email(email)
    if user is None:
        raise HTTPException(status_code=404, detail="User not found")

    user_id, _ = user
    chats = db_manager.get_chats_by_user_id(user_id)
    return {"chats": chats}


@app.delete("/delete_chat/{session_id}")
async def delete_chat(session_id: int):
    try:
        # Удаление чата из базы данных
        db_manager.delete_chat(session_id)
        return {"detail": "Chat deleted successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


def get_topic(session_id):
    try:
        with sqlite3.connect('metadata.db') as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT topic FROM Session WHERE id = ?", (session_id,))
            topic = cursor.fetchone()
        return topic[0] if topic else None
    except sqlite3.Error as e:
        print(f"Database error: {e}")
        return None


@app.get("/api/chat-topics")
async def get_chat_topics():
    conn = sqlite3.connect('metadata.db')
    cursor = conn.cursor()

    cursor.execute("SELECT id, topic FROM Session")
    rows = cursor.fetchall()

    conn.close()

    return [{"id": row[0], "topic": row[1]} for row in rows]


@app.websocket("/ws/rag_chat/")
async def websocket_endpoint(websocket: WebSocket):
    await websocket.accept()

    email = websocket.query_params.get('email')
    if email is None:
        await websocket.send_json({"error": "Требуется email"})
        await websocket.close(code=status.WS_1008_POLICY_VIOLATION)
        return

    user = get_user_by_email(email)
    if user is None:
        await websocket.send_json({"error": "Пользователь не найден"})
        await websocket.close(code=status.WS_1008_POLICY_VIOLATION)
        return

    user_id, cyberman_id = user

    session = get_session_by_user_and_cyberman(user_id, cyberman_id)
    if session is None:
        await websocket.send_json({"error": "Сессия не найдена"})
        await websocket.close(code=status.WS_1008_POLICY_VIOLATION)
        return

    session_id = session[0]
    topic = get_topic(session_id)

    chat_history_for_chain.current_session_id = session_id

    messages = db_manager.get_chat_messages_by_session_id(session_id)
    await websocket.send_json({"messages": messages, "topic": topic})
    try:
        while True:

            data = await websocket.receive_json()

            question_data = data.get('question_data')
            if question_data is None:
                await websocket.send_json({"error": "Требуется question_data"})
                continue

            question = question_data.get('question')
            new_session_id = question_data.get('session_id')  # Получаем session_id из сообщения
            if new_session_id:
                session_id = new_session_id  # Обновляем session_id, если он передан

            if question is None:
                await websocket.send_json({"error": "Требуется question"})
                continue

            try:
                answer = chain_with_message_history.invoke(
                    {"question": question, "context": format_docs(retriever.invoke(question))},
                    {"configurable": {"session_id": session_id}}
                ).content
            except Exception as e:
                await websocket.send_json({"error": str(e)})
                continue

            if answer:
                chat_history_for_chain.add_message(HumanMessage(content=question))
                chat_history_for_chain.add_message(AIMessage(content=answer))

            db_manager.add_chat_message(session_id, question, "human")  # Добавляем сообщение пользователя
            db_manager.add_chat_message(session_id, answer, "ai")  # Добавляем ответ бота

            # Получаем количество сообщений в сессии после добавления нового сообщения
            messages = db_manager.get_chat_messages_by_session_id(session_id)
            if len(messages) == 3:
                # Вызываем анализ текущей сессии только после третьего сообщения
                analyze_current_chat_topic(user_id, session_id)
                updated_topic = get_topic(session_id)
                await websocket.send_json({"topic_update": updated_topic})

            await websocket.send_json({"answer": answer})
    except WebSocketDisconnect:
        chat_history_for_chain.end_session()





def send_reset_email(email: str, token: str):
    sender_email = "datohar82@gmail.com"
    receiver_email = email
    password = "my_password"  # Не храните пароли в коде в реальных приложениях

    message = MIMEMultipart("alternative")
    message["Subject"] = "Подтверждение сброса пароля"
    message["From"] = sender_email
    message["To"] = receiver_email

    # reset_url = f"http://localhost:8222/confirm-reset?token={token}&email={email}"
    reset_url = f"https://chata100.up.railway.app/confirm-reset?token={token}&email={email}"

    text = f"Пожалуйста, подтвердите сброс пароля, перейдя по следующей ссылке: {reset_url}"
    html = f"""
    <html>
    <body>
        <p>Пожалуйста, подтвердите сброс пароля, перейдя по следующей ссылке:<br>
        <a href="{reset_url}">Подтвердить сброс пароля</a>
        </p>
    </body>
    </html>
    """

    part1 = MIMEText(text, "plain")
    part2 = MIMEText(html, "html")

    message.attach(part1)
    message.attach(part2)

    try:
        server = smtplib.SMTP_SSL("smtp.gmail.com", 465)
        server.login(sender_email, password)
        server.sendmail(sender_email, receiver_email, message.as_string())
        server.quit()
        print("Письмо успешно отправлено")
    except Exception as e:
        print(f"Ошибка при отправке письма: {e}")
        return JSONResponse(content={"status": "error", "message": "Пользователь не найден."}, status_code=404)



@app.get("/forgot-password", response_class=HTMLResponse)
async def get_forgot_password():
    return FileResponse("static/forgot_link.html")


# Сброс пароля
@app.post("/reset-password-request")
async def reset_password_request(email: str = Form(...)):
    with sqlite3.connect('metadata.db') as conn:
        cursor = conn.cursor()

        # Проверяем, существует ли пользователь с таким email
        cursor.execute("SELECT * FROM Users WHERE email = ?", (email,))
        user = cursor.fetchone()

        if user is None:
            return JSONResponse(content={"status": "error", "message": "Пользователь с таким email не найден."},
                                status_code=404)

        # Если пользователь найден, генерируем токен и обновляем запись
        token = get_password_hash(email)
        cursor.execute("UPDATE Users SET reset_token = ? WHERE email = ?", (token, email))
        conn.commit()

    send_reset_email(email, token)
    return JSONResponse(content={"status": "success"}, status_code=200)

#Обновление пароля
@app.post("/reset-password")
async def reset_password(hashed_password: str = Form(...), email: str = Form(...)):
    with sqlite3.connect('metadata.db') as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM Users WHERE email = ?", (email,))
        cursor.execute("UPDATE Users SET hashed_password = ? WHERE email = ?",
                       (get_password_hash(hashed_password), email))

        if cursor.rowcount == 0:
            return JSONResponse(content={"status": "error", "message": "Пользователь не найден."}, status_code=404)
        conn.commit()

        return JSONResponse(content={"status": "success", "redirect": "/login"})



# Шаг 3: Подтверждение сброса пароля по ссылке
@app.get("/confirm-reset")
async def confirm_reset(request: Request, token: str, email: str):
    with sqlite3.connect('metadata.db') as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT reset_token, new_password FROM Users WHERE email = ?", (email,))
        row = cursor.fetchone()
        if row and row[0] == token:
            cursor.execute(
                "UPDATE Users SET hashed_password = ?, reset_token = NULL, new_password = NULL WHERE email = ?",
                (row[1], email))
            conn.commit()
            # return RedirectResponse(url="/login")
            return FileResponse("static/forgot_password.html")
        else:
            return FileResponse("static/error_page.html")


from fastapi.templating import Jinja2Templates

templates = Jinja2Templates(directory="static")


def get_db():
    conn = sqlite3.connect('metadata.db', check_same_thread=False)
    try:
        yield conn.cursor()
    finally:
        conn.close()


def add_admin(email, password):
    conn = sqlite3.connect('metadata.db', check_same_thread=False)
    cursor = conn.cursor()
    hashed_password = pwd_context.hash(password)
    try:
        cursor.execute('''
        INSERT INTO Admin (email, password) VALUES (?, ?)
        ''', (email, hashed_password))
        conn.commit()
    except sqlite3.IntegrityError:
        print("Admin with this email already exists.")
    finally:
        conn.close()

# Добавление администратора
add_admin("datohar82@gmail.com", "UtlikA100")


@app.get("/admin/login", response_class=HTMLResponse)
async def admin_login_page(request: Request):
    return templates.TemplateResponse("admin_login.html", {"request": request})

@app.post("/admin/login")
async def admin_login(request: Request, email: str = Form(...), password: str = Form(...)):
    conn = sqlite3.connect('metadata.db', check_same_thread=False)
    cursor = conn.cursor()
    cursor.execute("SELECT password FROM Admin WHERE email = ?", (email,))
    admin = cursor.fetchone()
    conn.close()

    if admin and pwd_context.verify(password, admin[0]):
        request.session['admin'] = True
        return RedirectResponse(url="/admin", status_code=303)
    else:
        return templates.TemplateResponse("admin_login.html", {"request": request, "error": "Неверный email или пароль"})

@app.get("/admin", response_class=HTMLResponse)
async def admin_panel(request: Request):
    if not request.session.get('admin'):
        return RedirectResponse(url="/admin/login")

    conn = sqlite3.connect('metadata.db', check_same_thread=False)
    cursor = conn.cursor()
    cursor.execute("SELECT email, is_active FROM Users")
    users = cursor.fetchall()
    users_list = [{"email": user[0], "is_active": user[1]} for user in users]
    conn.close()

    return templates.TemplateResponse("admin_panel.html", {"request": request, "users": users_list})

@app.post("/admin/change-status")
async def change_status(request: Request, email: str = Form(...), current_status: str = Form(...)):
    if not request.session.get('admin'):
        return RedirectResponse(url="/admin/login")

    conn = sqlite3.connect('metadata.db', check_same_thread=False)
    cursor = conn.cursor()
    new_status = not bool(int(current_status))
    cursor.execute("UPDATE Users SET is_active = ? WHERE email = ?", (new_status, email))
    conn.commit()
    conn.close()

    return RedirectResponse(url="/admin", status_code=303)

@app.get("/admin/logout")
async def admin_logout(request: Request):
    request.session.clear()
    return RedirectResponse(url="/admin/login")




if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8222)
